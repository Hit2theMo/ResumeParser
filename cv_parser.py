import os
import re

# import PyPDF2
from docx import Document
# from PyPDF2 import PdfFileReader

# import nltk
# import pandas as pd
# import textract

# nltk.download('punkt', quiet=True)
# nltk.download('averaged_perceptron_tagger', quiet=True)
######################################################################
# USE THIS FUNCTION ONLY IF  file_To_Text() DOESNT WORK FOR PDF FILES


def pdf_To_Text(path):
    pdfFileObj = open(path, 'rb')
    pdfReader = PdfFileReader(pdfFileObj)
    pageObj = pdfReader.getPage(0)
    txt = pageObj.extractText()
    pdfFileObj.close()
    return txt

######################################################################


def file_To_Text(path):
    text = textract.process(path)
    return text.decode("utf-8")

# ----------------------------------------------------------------------------------------------------
# Function to extract the font names from PDF files


def extract_font_pdf(fname):
    def walk(obj, fnt, emb):
        if not hasattr(obj, 'keys'):
            return None, None
        fontkeys = set(['/FontFile', '/FontFile2', '/FontFile3'])
        if '/BaseFont' in obj:
            fnt.add(obj['/BaseFont'])
        if '/FontName' in obj:
            if [x for x in fontkeys if x in obj]:  # test to see if there is FontFile
                emb.add(obj['/FontName'])
        for k in obj.keys():
            walk(obj[k], fnt, emb)
        return fnt, emb

    pdf = PdfFileReader(fname)
    fonts = set()
    embedded = set()
    for page in pdf.pages:
        obj = page.getObject()
        if type(obj) == PyPDF2.generic.ArrayObject:
            for i in obj:
                if hasattr(i, 'keys'):
                    f, e = walk(i, fonts, embedded)
                    fonts = fonts.union(f)
                    embedded = embedded.union(e)
        else:
            f, e = walk(obj['/Resources'], fonts, embedded)
            fonts = fonts.union(f)
            embedded = embedded.union(e)

    font = sorted(list(fonts))
    for i in range(len(font)):
        font[i] = font[i].rstrip("MT")
        font[i] = font[i].lstrip("/")
    return font

# ----------------------------------------------------------------------------------------------------
# Function to extract the font name, font size, number of tables, images from a docx file


def extract_font_table_imgs_docx(path):
    doc = Document(path)
    font = []
    for p in doc.paragraphs:
        for r in p.runs:
            name = r.font.name
            size = r.font.size
            if size is not None:
                size = size / 12700
            if name is not None:
                if (name, size) not in font:
                    font.append((name, size))
    table_count = len(doc.tables)
    img_count = len(doc.inline_shapes)
    return font, table_count, img_count

# ----------------------------------------------------------------------------------------------------
# Function to extract email ids from both PDF and DOCX files


def extract_emails(txt):
    emails = re.findall(
        "([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)", txt)
    if emails == []:
        emails.append(None)
    return emails
# ----------------------------------------------------------------------------------------------------
# Function to extract mobile numbers from both PDF and DOCX files


def extract_mobile_number(text):
    phone = re.findall(re.compile(
        r'(?:(?:\+?([1-9]|[0-9][0-9]|[0-9][0-9][0-9])\s*(?:[.-]\s*)?)?(?:\(\s*([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9])\s*\)|([0-9][1-9]|[0-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9]))\s*(?:[.-]\s*)?)?([2-9]1[02-9]|[2-9][02-9]1|[2-9][02-9]{2})\s*(?:[.-]\s*)?([0-9]{4})(?:\s*(?:#|x\.?|ext\.?|extension)\s*(\d+))?'), text)
    if phone:
        number = ''.join(phone[0])
        if len(number) > 10:
            return '+' + number
        else:
            return number

# ----------------------------------------------------------------------------------------------------
# Function to extract the LinkedIn URL of user from both PDF and DOCX files


def extract_Linkedin(txt):
    url = re.search(
        r"http(s)?:\/\/([\w]+\.)?linkedin\.com\/in\/[A-z0-9_-]+\/?", txt)
    if url == None:
        return None
    else:
        return url.group()

# ----------------------------------------------------------------------------------------------------
# Function to extract the "INDIAN NAMES" from both PDF and DOCX files


def extract_name(document):

    # Reads Indian Names from the file, reduce all to lower case for easy comparision [Name lists]
    indianNames = open(r"indian_names.txt", "r").read().lower()
    indianNames = set(indianNames.split())
    otherNameHits = []
    nameHits = []
    name = None

    lines = [el.strip() for el in document.split("\n") if len(el)
             > 0]  # Splitting on the basis of newlines
    lines = [nltk.word_tokenize(el) for el in lines]
    lines = [nltk.pos_tag(el) for el in lines]

    sentences = nltk.sent_tokenize(document)
    # Split/Tokenize sentences into words (List of lists of strings)
    sentences = [nltk.word_tokenize(sent) for sent in sentences]
    tokens = sentences
    # Tag the tokens - list of lists of tuples - each tuple is (<word>, <tag>)
    sentences = [nltk.pos_tag(sent) for sent in sentences]
    dummy = []
    for el in tokens:
        dummy += el
    tokens = dummy
    # Try a regex chunk parser
    grammar = r'NAME: {<NN.*><NN.*><NN.*>*}'
    chunkParser = nltk.RegexpParser(grammar)
    all_chunked_tokens = []
    for tagged_tokens in lines:
        # Creates a parse tree
        if len(tagged_tokens) == 0:
            continue
        chunked_tokens = chunkParser.parse(tagged_tokens)
        all_chunked_tokens.append(chunked_tokens)
        for subtree in chunked_tokens.subtrees():
            if subtree.label() == 'NAME':
                for ind, leaf in enumerate(subtree.leaves()):
                    if leaf[0].lower() in indianNames and 'NN' in leaf[1]:
                        hit = " ".join([el[0]
                                        for el in subtree.leaves()[ind:ind + 3]])
                        if re.compile(r'[\d,:]').search(hit):
                            continue
                        nameHits.append(hit)
    if len(nameHits) > 0:
        nameHits = [re.sub(r'[^a-zA-Z \-]', '', el).strip() for el in nameHits]
        name = " ".join([el[0].upper() + el[1:].lower()
                         for el in nameHits[0].split() if len(el) > 0])
        otherNameHits = nameHits[1:]
    return name, otherNameHits

# ----------------------------------------------------------------------------------------------------
# Master function which combines all above methods to return a data frame


def extract_info(path):
    file_name, file_extension = os.path.splitext(path)
    if file_extension == ".pdf":
        txt = file_To_Text(path)
        font = extract_font_pdf(path)
        table_count = "NA"
        img_count = "NA"

    elif file_extension == ".docx":
        txt = file_To_Text(path)
        font, table_count, img_count = extract_font_table_imgs_docx(path)
    else:
        return "Invalid Format"

    linkedin = extract_Linkedin(txt)
    mobile = extract_mobile_number(txt)
    email = extract_emails(txt)
    name = extract_name(txt)[0]
    lines = len(txt.split("\n"))
    words = len(txt.split())
    chars = len(txt)

    data = [{
        "File Name": file_name + file_extension,
        "Name": name,
            "Contact Number": str(mobile),
            "Email ID(s)": str(email),
            "Linkedin URL": str(linkedin),
            "Total Lines": lines,
            "Total Characters": chars,
            "Total Words": words,
            "Fonts and Font sizes used": font,
            "Total number of Tables": table_count,
            "Total number of Images": img_count
            }]
    df = pd.DataFrame(data)
    return df


# ----------------------------------------------------------------------------------------------------
print("hello")
print(extract_font_table_imgs_docx("EY_Kitman Tsang_Cosec Mgr.docx"))
